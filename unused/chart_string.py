import os
import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_images(docx_path, image_folder):
    """提取图片并保存到指定文件夹，并返回图片位置字典"""
    doc = docx.Document(docx_path)
    image_positions = []
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img = rel.target_part.blob
            img_name = os.path.basename(rel.target_ref)  # 保持原始图片名称
            img_path = os.path.join(image_folder, img_name)
            # 保存图片
            with open(img_path, 'wb') as img_file:
                img_file.write(img)
            # 获取图片所在的段落位置
            for i, para in enumerate(doc.paragraphs):
                if rel.target_ref in para._element.xml:
                    image_positions.append((i, img_name))
                    break
    return image_positions

def extract_tables(docx_path, excel_folder):
    """提取表格并保存为Excel文件，并返回表格位置字典"""
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)
    doc = docx.Document(docx_path)
    table_positions = []
    for i, table in enumerate(doc.tables):
        df = pd.DataFrame(columns=[f'Column{j+1}' for j in range(len(table.columns))])
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            if len(row_data) != len(df.columns):
                print(f"Skipping row with mismatched columns: {row_data}")
                continue
            df.loc[len(df)] = row_data
        excel_path = os.path.join(excel_folder, f'table_{i+1}.xlsx')
        df.to_excel(excel_path, index=False)
        
        # 获取表格所在的段落位置
        for j, para in enumerate(doc.paragraphs):
            if table._element in para._element.getparent().getchildren():
                table_positions.append((j, f'table_{i+1}.xlsx'))
                break
    return table_positions

def extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions):
    """提取文本并转换为Markdown格式，并在图片和表格位置添加占位符"""
    doc = docx.Document(docx_path)
    markdown_text = ''
    image_counter = 1  # 图片计数器
    table_counter = 1  # 表格计数器
    
    # 遍历文档段落
    for i, para in enumerate(doc.paragraphs):
        # 处理段落内容
        if para.style and para.style.name.startswith('Heading'):
            level = para.style.name.split('Heading')[1] if len(para.style.name.split('Heading')) > 1 else "1"
            markdown_text += f'{"#" * int(level)} {para.text}\n\n'
        elif para.style and para.style.name == 'Normal':
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            markdown_text += f':::center\n{para.text}\n:::\n\n'

        # 处理文本格式
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f'**{text}**'
            if run.italic:
                text = f'*{text}*'
            if run.underline:
                text = f'_{text}_'
            if 'http' in run.text:
                text = f'[{text}]({text})'
            markdown_text += text

        # 检查是否需要在当前段落插入图片占位符
        for img_index, img_name in sorted(image_positions, key=lambda x: x[0]):
            if img_index == i:
                markdown_text += f'image{image_counter}\n\n'  # 使用 image1, image2 等占位符
                image_counter += 1
                image_positions.remove((img_index, img_name))  # 移除已处理的图片

        # 检查是否需要在当前段落插入表格占位符
        for table_index, table_file in sorted(table_positions, key=lambda x: x[0]):
            if table_index == i:
                markdown_text += f'table{table_counter}\n\n'  # 使用 table1, table2 等占位符
                table_counter += 1
                table_positions.remove((table_index, table_file))  # 移除已处理的表格

    # 保存Markdown文件
    markdown_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.md'
    markdown_path = os.path.join(markdown_folder, markdown_filename)
    if not os.path.exists(markdown_folder):
        os.makedirs(markdown_folder)
        
    with open(markdown_path, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_text)

def process_single_docx(docx_path, output_folder):
    """处理单个docx文件并生成Markdown文件，同时提取图片和表格"""
    # 获取文件名（不含扩展名），用于创建同名文件夹
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    doc_output_folder = os.path.join(output_folder, base_name)

    # 创建文件夹
    image_folder = os.path.join(doc_output_folder, 'images')
    excel_folder = os.path.join(doc_output_folder, 'excel')
    markdown_folder = os.path.join(doc_output_folder, 'markdown')

    # 提取图片和表格，并获取它们的位置
    image_positions = extract_images(docx_path, image_folder)
    table_positions = extract_tables(docx_path, excel_folder)

    # 生成Markdown文件
    extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions)


# 使用示例
docx_file_path = r"C:\Users\HP\Desktop\大创\output_step1"  # 输入的.docx文件路径
output_folder = r"C:\Users\HP\Desktop\大创\output_step2"  # 输出的文件夹路径

process_single_docx(docx_file_path, output_folder)
