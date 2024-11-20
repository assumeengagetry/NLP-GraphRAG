#这是使用LibreOffice的文件转化 doc to docx
import os
import subprocess

def convert_doc_to_docx(input_directory, output_directory, libreoffice_path):
    # 确保输出目录存在
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    
    # 遍历输入目录中的所有文件
    for filename in os.listdir(input_directory):
        if filename.endswith(".doc"):
            input_path = os.path.join(input_directory, filename)
            output_path = os.path.join(output_directory, os.path.splitext(filename)[0] + '.docx')
            
            # 构建 LibreOffice 命令
            command = [libreoffice_path, "--headless", "--convert-to", "docx", "--outdir", output_directory, input_path]
            
            # 调用 LibreOffice 进行转换
            subprocess.run(command, check=True)
            print(f"Converted {input_path} to {output_path}")

# 指定包含 .doc 文件的目录
input_directory = "C:\\Users\\HP\\Desktop\\大创\\input"
# 指定输出 .docx 文件的目录
output_directory = "C:\\Users\\HP\\Desktop\\大创\\output_step1"
# LibreOffice 的安装路径
libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.com"

convert_doc_to_docx(input_directory, output_directory, libreoffice_path)
import os
import shutil
import logging
from tqdm import tqdm
import pandas as pd
from docx import Document

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def move_docx_files(src_folder, dest_folder):
    for filename in os.listdir(src_folder):
        if filename.endswith(".docx"):
            src_file = os.path.join(src_folder, filename)
            dest_file = os.path.join(dest_folder, filename)
            shutil.move(src_file, dest_file)
            logging.info(f"Moved '{src_file}' to '{dest_file}'")

def extract_text(docx_file):
    doc = Document(docx_file)
    return '\n\n'.join([para.text for para in doc.paragraphs])

def extract_tables(docx_file):
    doc = Document(docx_file)
    tables_data = []
    for table in doc.tables:
        df = pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows])
        tables_data.append(df)
    return tables_data

def save_tables_to_excel(tables_data, excel_file):
    if not tables_data:  # 如果没有表格数据，不创建Excel文件
        logging.info(f"No tables to convert for {excel_file}")
        return
    with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
        for i, table_df in enumerate(tables_data):
            if not table_df.empty:  # 确保DataFrame不为空
                table_df.to_excel(writer, sheet_name=f'Table {i+1}', index=False)
            else:
                logging.info(f"Table {i+1} is empty, skipped.")
    logging.info(f"Tables saved to {excel_file}")

def excel_to_markdown(excel_file):
    xls = pd.ExcelFile(excel_file)
    markdown_content = ""
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        markdown_table = df.to_markdown(index=False)
        markdown_content += f"## {sheet_name}\n\n{markdown_table}\n\n"
    return markdown_content

def extract_images(docx_file, image_folder):
    doc = Document(docx_file)
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part.blob
            image_filename = os.path.basename(rel.target_ref)
            image_path = os.path.join(image_folder, image_filename)
            with open(image_path, 'wb') as img_file:
                img_file.write(image)
    logging.info(f"Images extracted to {image_folder}")

def process_docx_file(docx_file, output_md_folder, output_excel_folder, output_image_folder):
   # Ensure output directories exist
    if not os.path.exists(output_md_folder):
        os.makedirs(output_md_folder)
    if not os.path.exists(output_excel_folder):
        os.makedirs(output_excel_folder)
    if not os.path.exists(output_image_folder):
        os.makedirs(output_image_folder)

    # Extract text content
    markdown_content = extract_text(docx_file)
    
    # Extract tables data
    tables_data = extract_tables(docx_file)
    
    # Save tables as Excel file
    excel_file = os.path.splitext(os.path.basename(docx_file))[0] + '.xlsx'
    save_tables_to_excel(tables_data, os.path.join(output_excel_folder, excel_file))
    
    # Convert Excel to Markdown
    if os.path.exists(os.path.join(output_excel_folder, excel_file)):
        markdown_tables_content = excel_to_markdown(os.path.join(output_excel_folder, excel_file))
    else:
        markdown_tables_content = ""
    
    # Combine text and tables Markdown content
    markdown_content_with_tables = markdown_content + "\n" + markdown_tables_content
    
    # Save Markdown file
    markdown_file = os.path.splitext(os.path.basename(docx_file))[0] + '.md'
    with open(os.path.join(output_md_folder, markdown_file), 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_content_with_tables)
    logging.info(f"Markdown file saved: {os.path.join(output_md_folder, markdown_file)}")
    
    # Extract images
    extract_images(docx_file, output_image_folder)
def main():
    src_folder = r"C:\Users\HP\Desktop\大创\input"
    dest_folder = r"C:\Users\HP\Desktop\大创\output_step1"
    
    move_docx_files(src_folder, dest_folder)
    
    input_folder = dest_folder
    output_md_folder = r"C:\Users\HP\Desktop\大创\output_step2"
    output_excel_folder = r"C:\Users\HP\Desktop\大创\chart_excel"
    output_image_folder = r"C:\Users\HP\Desktop\大创\pictures"
    
    for filename in tqdm(os.listdir(input_folder), desc="Processing files"):
        if filename.endswith(".docx"):
            docx_file_path = os.path.join(input_folder, filename)
            process_docx_file(docx_file_path, output_md_folder, output_excel_folder, output_image_folder)

if __name__ == "__main__":
    main()