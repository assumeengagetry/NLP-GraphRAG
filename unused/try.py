import os
import docx
from docx.oxml.ns import qn

def insert_image_placeholders(docx_path, output_path):
    doc = docx.Document(docx_path)
    image_counter = 1

    # 遍历所有段落和运行（run），并检查其中是否包含图片
    for para in doc.paragraphs:
        for run in para.runs:
            # 获取运行中的元素
            inline_shapes = run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill")
            if inline_shapes:
                # 如果找到了图像，则插入占位符
                placeholder = f"@@image{image_counter}@@"
                para.add_run(placeholder)  # 在段落中插入占位符
                image_counter += 1

    # 保存修改后的文档
    doc.save(output_path)
    print(f"Document with image placeholders saved to {output_path}")

# 示例用法
input_path = r"C:\Users\HP\Desktop\大创\output_step1\3-正宏商混-现评-正式版.docx"  # 输入的 .docx 文件路径
output_path = r"C:\Users\HP\Desktop\大创\output_step1\3-正宏商混-现评-修改版.docx"  # 输出的 .docx 文件路径

insert_image_placeholders(input_path, output_path)
