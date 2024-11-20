import os
from docx import Document

def insert_table_placeholders(docx_path, output_path):
    """在表格的位置插入@@table1@@, @@table2@@等占位符"""
    doc = Document(docx_path)
    table_counter = 1  # 用于计数表格

    # 遍历文档中的所有表格
    for table in doc.tables:
        # 在表格前插入占位符
        table_placeholder = f"@@table{table_counter}@@"
        paragraph = doc.add_paragraph(table_placeholder)
        # 将段落移动到表格前的位置
        table._element.addprevious(paragraph._element)
        table_counter += 1

    # 保存修改后的文档
    doc.save(output_path)
    print(f"Document with table placeholders saved to {output_path}")

# 使用示例
input_path = r"C:\Users\HP\Desktop\大创\output_step1\3-正宏商混-现评-修改版.docx"  # .docx 文件路径
output_path = r"C:\Users\HP\Desktop\大创\output_step1\3-正宏商混-现评-再修改版.docx"  # 输出的 .docx 文件路径

insert_table_placeholders(input_path, output_path)