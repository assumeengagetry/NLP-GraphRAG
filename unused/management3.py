import os
import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_placeholder(para, placeholder):
    """在段落中插入占位符"""
    para.add_run(f'\n{placeholder}\n')  # 在段落中插入占位符

def extract_images(docx_path, image_folder):
    """提取图片并保存到指定文件夹，并返回图片位置字典"""
    doc = docx.Document(docx_path)
    image_positions = []
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    image_count = 1  # 用于给图片编号
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img = rel.target_part.blob
            img_name = os.path.basename(rel.target_ref)  # 保持原始图片名称
            img_path = os.path.join(image_folder, img_name)
            # 保存图片
            with open(img_path, 'wb') as img_file:
                img_file.write(img)
            # 获取图片所在的段落位置
            for i, para in enumerate(doc.paragraphs):
                if rel.target_ref in para._element.xml:
                    placeholder = f'![image{image_count}]({img_name})'  # 使用占位符格式
                    insert_placeholder(para, placeholder)  # 插入占位符
                    image_positions.append((i, placeholder))  # 记录占位符的位置
                    image_count += 1
                    break
    return image_positions

def extract_tables(docx_path, excel_folder):
    """提取表格并保存为Excel文件，并返回表格位置字典"""
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)
    doc = docx.Document(docx_path)
    table_positions = []
    table_count = 1  # 用于给表格编号
    for i, table in enumerate(doc.tables):
        df = pd.DataFrame(columns=[f'Column{j+1}' for j in range(len(table.columns))])
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            if len(row_data) != len(df.columns):
                print(f"Skipping row with mismatched columns: {row_data}")
                continue
            df.loc[len(df)] = row_data
        excel_path = os.path.join(excel_folder, f'table_{table_count}.xlsx')
        df.to_excel(excel_path, index=False)
        
        # 获取表格所在的段落位置
        for j, para in enumerate(doc.paragraphs):
            if table._element in para._element.getparent().getchildren():
                placeholder = f'table_{table_count}'  # 占位符格式
                insert_placeholder(para, placeholder)  # 插入占位符
                table_positions.append((j, placeholder))  # 记录占位符的位置
                table_count += 1
                break
    return table_positions

def extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions):
    """提取文本并转换为Markdown格式，并在图片和表格位置添加占位符"""
    doc = docx.Document(docx_path)
    markdown_text = ''
    
    # 遍历文档段落
    for i, para in enumerate(doc.paragraphs):
        # 检查是否需要在当前段落插入图片占位符
        for img_index, img_placeholder in sorted(image_positions, key=lambda x: x[0]):
            if img_index == i:
                markdown_text += f'![{img_placeholder}]()  # Placeholder for {img_placeholder}\n\n'
                image_positions.remove((img_index, img_placeholder))  # 移除已处理的图片

        # 检查是否需要在当前段落插入表格占位符
        for table_index, table_placeholder in sorted(table_positions, key=lambda x: x[0]):
            if table_index == i:
                markdown_text += f'\n\n{table_placeholder}  # Placeholder for {table_placeholder}\n\n'
                table_positions.remove((table_index, table_placeholder))  # 移除已处理的表格

        try:
            # 处理段落内容
            if para.style and para.style.name.startswith('Heading'):
                level = para.style.name.split('Heading')[1] if len(para.style.name.split('Heading')) > 1 else "1"
                markdown_text += f'{"#" * int(level)} {para.text}\n\n'
            elif para.style and para.style.name == 'Normal':
                markdown_text += f'{para.text}\n\n'
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                markdown_text += f'{para.text}\n\n'
            elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                markdown_text += f':::center\n{para.text}\n:::\n\n'
            
            # 处理文本格式
            for run in para.runs:
                text = run.text
                if run.bold:
                    text = f'**{text}**'
                if run.italic:
                    text = f'*{text}*'
                if run.underline:
                    text = f'_{text}_'
                if 'http' in run.text:
                    text = f'[{text}]({text})'
                markdown_text += text
        except Exception as e:
            print(f"Error processing paragraph {i}: {e}")

    # 保存Markdown文件
    markdown_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.md'
    markdown_path = os.path.join(markdown_folder, markdown_filename)
    with open(markdown_path, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_text)

def process_docx_folder(docx_folder, output_folder):
    """处理文件夹中的所有docx文件并生成Markdown文件"""
    image_folder = os.path.join(output_folder, 'images')
    excel_folder = os.path.join(output_folder, 'excel')
    markdown_folder = os.path.join(output_folder, 'markdown')
    
    if not os.path.exists(markdown_folder):
        os.makedirs(markdown_folder)
    
    for filename in os.listdir(docx_folder):
        if filename.endswith('.docx'):
            docx_path = os.path.join(docx_folder, filename)
            # 提取图片和表格，并获取它们的位置
            image_positions = extract_images(docx_path, image_folder)
            table_positions = extract_tables(docx_path, excel_folder)
            # 生成Markdown文件
            extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions)

# 使用示例
docx_folder = r"C:\Users\HP\Desktop\大创\output_step1"  # 输入的.docx文件夹路径
output_folder = r"C:\Users\HP\Desktop\大创\output_step2"  # 输出的Markdown文件夹路径

process_docx_folder(docx_folder, output_folder)
