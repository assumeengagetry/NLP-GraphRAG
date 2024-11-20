import os
import subprocess
import docx
import pandas as pd
import re

# 使用 LibreOffice 将 .doc 转为 .docx
def convert_doc_to_docx(input_directory, libreoffice_path):
    for filename in os.listdir(input_directory):
        if filename.endswith(".doc"):
            input_path = os.path.join(input_directory, filename)
            output_path = os.path.join(input_directory, os.path.splitext(filename)[0] + '.docx')
            command = [libreoffice_path, "--headless", "--convert-to", "docx", "--outdir", input_directory, input_path]
            subprocess.run(command, check=True)
            print(f"Converted {input_path} to {output_path}")

# 提取图片并保存到指定文件夹，并返回图片位置字典
def extract_images(docx_path, image_folder):
    doc = docx.Document(docx_path)
    image_positions = []
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    for rel in doc.part.rels.values():
        if r"C:\Users\HP\Desktop\大创\pictures" in rel.reltype:
            try:
                img = rel.target_part.blob
                img_name = os.path.basename(rel.target_ref)
                img_path = os.path.join(image_folder, img_name)
                with open(img_path, 'wb') as img_file:
                    img_file.write(img)
                for i, para in enumerate(doc.paragraphs):
                    if rel.target_ref in para._element.xml:
                        image_positions.append((i, img_name, img_path))
                        break
            except ValueError as e:
                print(f"Error extracting image {rel.target_ref}: {e}")
    
    return image_positions

# 提取表格并保存为Excel文件，并返回表格位置字典
def extract_tables(docx_path, excel_folder):
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)
    doc = docx.Document(docx_path)
    table_positions = []
    for i, table in enumerate(doc.tables):
        df = pd.DataFrame(columns=[f'Column{j+1}' for j in range(len(table.columns))])
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            if len(row_data) != len(df.columns):
                print(f"Skipping row with mismatched columns: {row_data}")
                continue
            df.loc[len(df)] = row_data
        excel_path = os.path.join(excel_folder, f'table_{i+1}.xlsx')
        df.to_excel(excel_path, index=False)
        
        for j, para in enumerate(doc.paragraphs):
            if table._element in para._element.getparent().getchildren():
                table_positions.append((j, f'table_{i+1}.xlsx', excel_path))
                break
    return table_positions

# 将Excel表格转换为Markdown格式
def convert_excel_to_markdown(excel_path):
    df = pd.read_excel(excel_path)
    return df.to_markdown(index=False)

# 提取文本并转换为Markdown格式
def extract_text_to_markdown(docx_path, image_folder, excel_folder, image_positions, table_positions):
    doc = docx.Document(docx_path)
    markdown_text = ''
    
    for i, para in enumerate(doc.paragraphs):
        for img_index, img_name, img_path in sorted(image_positions, key=lambda x: x[0]):
            if img_index == i:
                markdown_text += f'![{img_name}]({img_path})\n\n'
                image_positions.remove((img_index, img_name, img_path))
        
        for table_index, table_file, excel_path in sorted(table_positions, key=lambda x: x[0]):
            if table_index == i:
                table_md = convert_excel_to_markdown(excel_path)
                markdown_text += f'\n\n{table_md}\n\n'
                table_positions.remove((table_index, table_file, excel_path))

        # 处理文本格式
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f'**{text}**'
            if run.italic:
                text = f'*{text}*'
            if run.underline:
                text = f'_{text}_'
            if 'http' in run.text:
                text = f'[{text}]({text})'
            markdown_text += text
        
        markdown_text += '\n\n'  # 段落之间添加换行
    
    return markdown_text


# 按指定格式的标题切分Markdown内容并存入数组
def split_markdown_by_custom_headers(markdown_text):
    split_content_array = []
    # 更新正则表达式，只匹配 **1.1** 形式的小标题
    header_pattern = re.compile(r'\*\*(\d+\.\d+)\*\*\s*(.*?)\s*(?=\*\*|\Z)')
    matches = list(header_pattern.finditer(markdown_text))
    
    if not matches:
        print("未找到任何匹配的标题")
        return split_content_array
    
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(markdown_text)
        header_number = matches[i].group(1).strip()
        header_text = matches[i].group(2).strip()
        section_content = markdown_text[start:end].strip()
        
        split_content_array.append({
            'header_number': header_number,
            'header_text': header_text,
            'content': section_content
        })
    
    return split_content_array

# 将分节内容生成 Markdown 文件
def generate_markdown_file(split_content_array, output_directory):
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    output_file_path = os.path.join(output_directory, 'split_content.md')
    
    with open(output_file_path, 'w', encoding='utf-8') as md_file:
        for section in split_content_array:
            md_file.write(f"**{section['header_number']}** {section['header_text']}\n\n")
            md_file.write(f"{section['content']}\n\n")
    
    print(f"切分后的 Markdown 文件已生成：{output_file_path}")

# 处理文件夹中的所有.docx文件
def process_docx_folder(input_directory, libreoffice_path):
    # Step 1: 将 .doc 文件转换为 .docx
    convert_doc_to_docx(input_directory, libreoffice_path)
    
    split_content_array = []
    image_folder = os.path.join(input_directory, 'images')
    excel_folder = os.path.join(input_directory, 'excel')
    
    for filename in os.listdir(input_directory):
        if filename.endswith('.docx'):
            docx_path = os.path.join(input_directory, filename)
            image_positions = extract_images(docx_path, image_folder)
            table_positions = extract_tables(docx_path, excel_folder)
            markdown_text = extract_text_to_markdown(docx_path, image_folder, excel_folder, image_positions, table_positions)
            
            # 打印提取到的 Markdown 内容
            print("提取到的 Markdown 内容：")
            print(markdown_text)
            
            # Step 4: 按标题切分Markdown内容并存入数组
            split_content = split_markdown_by_custom_headers(markdown_text)
            split_content_array.extend(split_content)
    
    # 生成 Markdown 文件
    output_directory = os.path.join(input_directory, 'output')  # 指定输出文件夹
    generate_markdown_file(split_content_array, output_directory)

# 使用示例
input_directory = r"C:\Users\HP\Desktop\大创\input"  # 输入文件夹路径
libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.com"  # LibreOffice 的安装路径

process_docx_folder(input_directory, libreoffice_path)
