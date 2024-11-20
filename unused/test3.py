import os
import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_images(docx_path, image_folder):
    """提取图片并保存到指定文件夹，并返回图片位置字典"""
    doc = docx.Document(docx_path)
    image_positions = []
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img = rel.target_part.blob
            img_name = os.path.basename(rel.target_ref)
            img_path = os.path.join(image_folder, img_name)
            # 保存图片
            with open(img_path, 'wb') as img_file:
                img_file.write(img)
            # 获取图片所在的段落位置
            for i, para in enumerate(doc.paragraphs):
                if rel.target_ref in para._element.xml:
                    image_positions.append((i, img_name))
                    break
    return image_positions

def extract_tables(docx_path, excel_folder):
    """提取表格并保存为Excel文件，并返回表格位置字典"""
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)
    doc = docx.Document(docx_path)
    table_positions = []
    for i, table in enumerate(doc.tables):
        df = pd.DataFrame(columns=[f'Column{j+1}' for j in range(len(table.columns))])
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            if len(row_data) != len(df.columns):
                print(f"Skipping row with mismatched columns: {row_data}")
                continue
            df.loc[len(df)] = row_data
        excel_path = os.path.join(excel_folder, f'table_{i+1}.xlsx')
        df.to_excel(excel_path, index=False)
        
        # 获取表格所在的段落位置
        for j, para in enumerate(doc.paragraphs):
            if table._element in para._element.getparent().getchildren():
                table_positions.append((j, f'table_{i+1}.xlsx'))
                break
    return table_positions

def convert_excel_to_markdown(excel_path):
    """将Excel表格转换为Markdown格式"""
    df = pd.read_excel(excel_path)
    return df.to_markdown(index=False)

def extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions):
    """提取文本并转换为Markdown格式，并在图片和表格位置添加占位符"""
    if not os.path.exists(markdown_folder):
        os.makedirs(markdown_folder)
    doc = docx.Document(docx_path)
    markdown_text = ''
    
    # 遍历文档段落
    for i, para in enumerate(doc.paragraphs):
        # 如果当前段落中应该有图片占位符
        for img_index, img_name in image_positions:
            if img_index == i:
                img_path = os.path.join(image_folder, img_name)
                markdown_text += f'![{img_name} Placeholder]({img_path})\n\n'
        
        # 如果当前段落中应该有表格占位符
        for table_index, table_file in table_positions:
            if table_index == i:
                excel_path = os.path.join(excel_folder, table_file)
                table_md = convert_excel_to_markdown(excel_path)
                markdown_text += f'\n\n{table_md}\n\n'
        
        # 处理段落内容
        if para.style and para.style.name.startswith('Heading'):
            level = para.style.name.split('Heading')[1] if len(para.style.name.split('Heading')) > 1 else "1"
            markdown_text += f'{"#" * int(level)} {para.text}\n\n'
        elif para.style and para.style.name == 'Normal':
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            markdown_text += f':::center\n{para.text}\n:::\n\n'
        
        # 处理文本格式
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f'**{text}**'
            if run.italic:
                text = f'*{text}*'
            if run.underline:
                text = f'_{text}_'
            if 'http' in run.text:
                text = f'[{text}]({text})'
            markdown_text += text

    markdown_path = os.path.join(markdown_folder, 'document.md')
    with open(markdown_path, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_text)

# 使用示例
docx_path = r"C:\Users\HP\Desktop\大创\output_step1\3-正宏商混-现评-正式版.docx"
image_folder = r"C:\Users\HP\Desktop\大创\pictures"
excel_folder = r"C:\Users\HP\Desktop\大创\chart_excel"
markdown_folder = r"C:\Users\HP\Desktop\大创\output_step2_md"

# 提取图片和表格，并获取它们的位置
image_positions = extract_images(docx_path, image_folder)
table_positions = extract_tables(docx_path, excel_folder)

# 生成Markdown文件
extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions)
