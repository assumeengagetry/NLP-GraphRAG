from docx import Document
import pandas as pd
import os

def extract_content(docx_file):
    doc = Document(docx_file)
    markdown_content = []

    # Extract text
    for para in doc.paragraphs:
        markdown_content.append(para.text + '\n\n')

    # Extract tables
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append('| ' + ' | '.join(row_content) + ' |')
        markdown_content.append('\n'.join(table_content) + '\n\n')

    return ''.join(markdown_content)

def extract_tables(docx_file):
    doc = Document(docx_file)
    tables_data = []
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append(row_content)
        tables_data.append(table_content)
    return tables_data

def save_tables_to_excel(tables_data, output_file):
    with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
        for i, table in enumerate(tables_data):
            df = pd.DataFrame(table)
            df.to_excel(writer, sheet_name=f'Table {i+1}', index=False)
        
        # Ensure at least one sheet is visible
        workbook = writer.book
        if all(sheet.sheet_state == 'hidden' for sheet in workbook.worksheets):
            workbook.active.sheet_state = 'visible'

def excel_to_markdown(excel_file):
    xls = pd.ExcelFile(excel_file)
    markdown_content = ""
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        markdown_table = df.to_markdown(index=False)
        markdown_content += f"## {sheet_name}\n\n{markdown_table}\n\n"
    return markdown_content

def extract_images(docx_file, image_folder):
    doc = Document(docx_file)
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part.blob
            image_filename = os.path.basename(rel.target_ref)
            image_path = os.path.join(image_folder, image_filename)
            with open(image_path, 'wb') as img_file:
                img_file.write(image)
    print(f"Images extracted to {image_folder}")

def create_image_markdown(image_folder, markdown_file):
    image_filenames = [f for f in os.listdir(image_folder) if os.path.isfile(os.path.join(image_folder, f))]
    markdown_content = ""
    for image_filename in image_filenames:
        image_path = os.path.join(image_folder, image_filename)
        markdown_content += f"![]({image_path})\n\n"
    with open(markdown_file, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_content)
    print(f"Markdown file with images created at {markdown_file}")

def main(input_folder, output_markdown_folder, image_folder, excel_folder):
    # Ensure output folders exist
    os.makedirs(output_markdown_folder, exist_ok=True)
    os.makedirs(image_folder, exist_ok=True)
    os.makedirs(excel_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        if filename.endswith('.docx'):
            docx_file_path = os.path.join(input_folder, filename)

            # Extract content
            markdown_content = extract_content(docx_file_path)
            markdown_file_path = os.path.join(output_markdown_folder, f'{os.path.splitext(filename)[0]}.md')
            with open(markdown_file_path, 'w', encoding='utf-8') as file:
                file.write(markdown_content)

            # Extract tables and save to Excel
            tables_data = extract_tables(docx_file_path)
            excel_file_path = os.path.join(excel_folder, f'{os.path.splitext(filename)[0]}_tables.xlsx')
            save_tables_to_excel(tables_data, excel_file_path)

            # Extract images
            extract_images(docx_file_path, image_folder)

            # Create image markdown
            image_markdown_file_path = os.path.join(output_markdown_folder, f'{os.path.splitext(filename)[0]}_images.md')
            create_image_markdown(image_folder, image_markdown_file_path)

if __name__ == "__main__":
    input_folder = r"C:\Users\HP\Desktop\大创\output_step1"
    output_markdown_folder = r"C:\Users\HP\Desktop\大创\output_step2"
    image_folder = r"C:\Users\HP\Desktop\大创\pictures_child"
    excel_folder = r"C:\Users\HP\Desktop\大创\chart_child"
    main(input_folder, output_markdown_folder, image_folder, excel_folder)
