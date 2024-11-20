import os
import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_folder(folder_path):
    """创建文件夹，如果已存在则忽略"""
    os.makedirs(folder_path, exist_ok=True)

def extract_images(docx_path, image_folder):
    """提取图片并保存到指定文件夹，忽略外部链接"""
    doc = docx.Document(docx_path)
    image_counter = 1

    create_folder(image_folder)

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            if rel.target_ref.startswith('http'):
                print(f"跳过外部图片: {rel.target_ref}")
                continue

            # 获取图片扩展名
            img_extension = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
            img_name = f'image{image_counter}.{img_extension}'
            img_path = os.path.join(image_folder, img_name)
            
            try:
                img = rel.target_part.blob
                with open(img_path, 'wb') as img_file:
                    img_file.write(img)
                image_counter += 1
            except Exception as e:
                print(f"保存图片 {rel.target_ref} 时出错: {e}")

def extract_tables(docx_path, excel_folder):
    """提取表格并保存为Excel文件"""
    create_folder(excel_folder)
    doc = docx.Document(docx_path)
    table_counter = 1
    
    for table in doc.tables:
        # 获取第一行作为列名
        headers = [cell.text for cell in table.rows[0].cells]
        df = pd.DataFrame(columns=headers)  # 定义列名

        for row in table.rows[1:]:  # 从第二行开始读取数据
            row_data = [cell.text for cell in row.cells]
            df.loc[len(df)] = row_data

        excel_path = os.path.join(excel_folder, f'table{table_counter}.xlsx')
        df.to_excel(excel_path, index=False)
        table_counter += 1

def extract_text_to_markdown(docx_path, markdown_folder):
    """提取文本并转换为Markdown格式"""
    doc = docx.Document(docx_path)
    markdown_text = ''
    
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = para.style.name.split('Heading')[1] if len(para.style.name.split('Heading')) > 1 else "1"
            markdown_text += f'{"#" * int(level)} {para.text}\n\n'
        elif para.style.name == 'Normal':
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            markdown_text += f':::center\n{para.text}\n:::\n\n'
        
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f'**{text}** '
            if run.italic:
                text = f'*{text}*'
            if run.underline:
                text = f'_{text}_'
            if re.match(r'http[s]?://', text):  # 正则表达式匹配 URL
                text = f'[{text}]({text})'
            markdown_text += text

    markdown_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.md'
    markdown_path = os.path.join(markdown_folder, markdown_filename)
    create_folder(markdown_folder)
    with open(markdown_path, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_text)

def process_single_docx(docx_path, output_folder):
    """处理单个docx文件并生成Markdown文件"""
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    image_folder = os.path.join(output_folder, base_name, 'images')
    excel_folder = os.path.join(output_folder, base_name, 'excel')
    markdown_folder = os.path.join(output_folder, base_name, 'markdown')
    
    create_folder(image_folder)
    create_folder(excel_folder)
    create_folder(markdown_folder)

    extract_images(docx_path, image_folder)
    extract_tables(docx_path, excel_folder)
    extract_text_to_markdown(docx_path, markdown_folder)

def process_docx_folder(docx_folder, output_folder):
    """处理文件夹中的所有docx文件并生成Markdown文件"""
    create_folder(output_folder)
    for filename in os.listdir(docx_folder):
        if filename.endswith('.docx'):
            docx_path = os.path.join(docx_folder, filename)
            process_single_docx(docx_path, output_folder)

# 使用示例
docx_folder = r"C:\Users\HP\Desktop\大创\output_step1"  # 输入的 .docx 文件夹路径
output_folder = r"C:\Users\HP\Desktop\大创\output_step2"  # 输出的 Markdown 文件夹路径

process_docx_folder(docx_folder, output_folder)
