import os
import docx
import openpyxl
from docx import Document
from openpyxl import Workbook

def extract_images(docx_path, image_folder):
    """提取图片并保存到指定文件夹，并返回图片位置字典"""
    doc = Document(docx_path)
    image_positions = []
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img = rel.target_part.blob
            img_name = os.path.basename(rel.target_ref)
            img_path = os.path.join(image_folder, img_name)
            with open(img_path, 'wb') as img_file:
                img_file.write(img)
            image_positions.append(img_path)
    return image_positions

def extract_tables(docx_path, excel_folder, excel_name="tables.xlsx"):
    """提取表格并将其保存为Excel文件，返回Excel文件路径"""
    doc = Document(docx_path)
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)
    
    # 创建Excel文件
    excel_path = os.path.join(excel_folder, excel_name)
    wb = Workbook()
    
    for table_index, table in enumerate(doc.tables):
        ws = wb.create_sheet(title=f"Table_{table_index + 1}")
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                ws.cell(row=row_index + 1, column=col_index + 1).value = cell.text
    
    wb.save(excel_path)
    return excel_path

def main(docx_path, output_folder):
    # 创建输出文件夹
    image_folder = os.path.join(output_folder, "images")
    excel_folder = os.path.join(output_folder, "tables")
    
    # 提取图片
    image_paths = extract_images(docx_path, image_folder)
    print(f"图片已保存到文件夹：{image_folder}")
    
    # 提取表格
    excel_path = extract_tables(docx_path, excel_folder)
    print(f"表格已保存为Excel文件：{excel_path}")

# 示例调用
docx_path = r"C:\Users\HP\Desktop\大创\output_step1\06睿恒化工--控评--修改版2.docx"
output_folder = r"C:\Users\HP\Desktop\大创\output_step2"
main(docx_path, output_folder)
