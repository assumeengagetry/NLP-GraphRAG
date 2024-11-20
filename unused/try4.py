import os
import docx
from docx.oxml.ns import qn

def insert_image_placeholders(docx_path, output_path):
    doc = docx.Document(docx_path)
    image_counter = 1

    # 遍历所有段落和运行（run），并检查其中是否包含图片
    for para in doc.paragraphs:
        for run in para.runs:
            # 获取运行中的元素
            inline_shapes = run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill")
            if inline_shapes:
                # 如果找到了图像，则插入占位符
                placeholder = f"\n@@image{image_counter}@@"
                para.add_run(placeholder)  # 在段落中插入占位符
                image_counter += 1

    # 保存修改后的文档
    doc.save(output_path)
    print(f"Document with image placeholders saved to {output_path}")

# 示例用法
input_path = r"C:\Users\HP\Desktop\大创\output_step1\06睿恒化工--控评--正式版.docx"  # 输入的 .docx 文件路径
output_path = r"C:\Users\HP\Desktop\大创\output_step1\06睿恒化工--控评--修改版1.docx"  # 输出的 .docx 文件路径

insert_image_placeholders(input_path, output_path)
import os
from docx import Document

def insert_table_placeholders(docx_path, output_path):
    """在表格的位置插入@@table1@@, @@table2@@等占位符"""
    doc = Document(docx_path)
    table_counter = 1  # 用于计数表格

    # 遍历文档中的所有表格
    for table in doc.tables:
        # 在表格前插入占位符
        table_placeholder = f"@@table{table_counter}@@"
        paragraph = doc.add_paragraph(table_placeholder)
        # 将段落移动到表格前的位置
        table._element.addprevious(paragraph._element)
        table_counter += 1

    # 保存修改后的文档
    doc.save(output_path)
    print(f"Document with table placeholders saved to {output_path}")

# 使用示例
input_path = r"C:\Users\HP\Desktop\大创\output_step1\06睿恒化工--控评--修改版1.docx"# .docx 文件路径
output_path = r"C:\Users\HP\Desktop\大创\output_step1\06睿恒化工--控评--修改版2.docx"  # 输出的 .docx 文件路径

insert_table_placeholders(input_path, output_path)