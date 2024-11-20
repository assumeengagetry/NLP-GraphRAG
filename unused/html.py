import os
from docx import Document
from PIL import Image
from io import BytesIO

def convert_docx_to_html(docx_path, html_path):
    doc = Document(docx_path)
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(html_path), exist_ok=True)

    with open(html_path, "w", encoding="utf-8") as html_file:
        # 写入 HTML 头部
        html_file.write("<html>\n<head>\n<title>{}</title>\n</head>\n<body>\n".format(doc.core_properties.title if doc.core_properties.title else "Untitled Document"))
        
        # 遍历文档中的每个段落
        for para in doc.paragraphs:
            html_file.write("<p>{}</p>\n".format(para.text))

        # 遍历文档中的每个表格
        for table in doc.tables:
            html_file.write("<table border='1'>\n")
            for row in table.rows:
                html_file.write("<tr>\n")
                for cell in row.cells:
                    html_file.write("<td>{}</td>\n".format(cell.text))
                html_file.write("</tr>\n")
            html_file.write("</table>\n")
        
        # 遍历文档中的每个图像
        image_counter = 1
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image = rel.target_part.blob
                image_stream = BytesIO(image)
                image_format = rel.target_ref.format
                image_filename = f"image{image_counter}.{image_format}"
                image_path = os.path.join(os.path.dirname(html_path), image_filename)
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_stream.read())
                html_file.write(f"<img src='{image_filename}' alt='Image {image_counter}' />\n")
                image_counter += 1
        
        # 写入 HTML 尾部
        html_file.write("</body>\n</html>")

    print(f"Converted {docx_path} to {html_path}")

# 示例用法
input_path = r"C:\Users\HP\Desktop\大创\output_step1\3-正宏商混-现评-正式版.docx"  # .docx 文件路径
output_path = r"C:\Users\HP\Desktop\大创\output_step2\example.html"  # 输出的 .html 文件路径
convert_docx_to_html(input_path, output_path)