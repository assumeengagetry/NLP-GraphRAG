import re
from docx import Document
import os

def get_heading_level(style_name):
    # 使用正则表达式匹配数字部分
    match = re.search(r'\d+', style_name)
    if match:
        return int(match.group())
    return float('inf')  # 如果没有找到数字，返回一个很大的数，表示这是一个普通段落

def split_docx_by_headings(doc_path, output_dir):
    doc = Document(doc_path)
    
    current_doc = None
    current_heading = None
    part_number = 1

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            if current_doc:
                # 保存当前文档
                output_path = os.path.join(output_dir, f"{part_number}_{current_heading}.docx")
                current_doc.save(output_path)
                part_number += 1

            # 创建新文档并添加标题
            current_doc = Document()
            current_heading = paragraph.text
            current_doc.add_paragraph(paragraph.text, style=paragraph.style.name)
        elif current_doc:
            # 添加非标题段落到当前文档
            current_doc.add_paragraph(paragraph.text, style='Normal')

    # 保存最后一个文档
    if current_doc:
        output_path = os.path.join(output_dir, f"{part_number}_{current_heading}.docx")
        current_doc.save(output_path)

if __name__ == "__main__":
    input_path = r"C:\Users\HP\Desktop\大创\output_step1\06睿恒化工--控评--修改版2.docx"# 替换为你的输入文件路径
    output_directory = r"C:\Users\HP\Desktop\大创\output_step2"  # 替换为你的输出目录路径
    split_docx_by_headings(input_path, output_directory)