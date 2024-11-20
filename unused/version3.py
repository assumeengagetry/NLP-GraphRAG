import os
import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_images(docx_path, image_folder):
    """提取图片并保存到指定文件夹，并返回图片位置字典"""
    doc = docx.Document(docx_path)
    image_positions = []
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img = rel.target_part.blob
            img_name = os.path.basename(rel.target_ref)  # 保持原始图片名称
            img_path = os.path.join(image_folder, img_name)
            # 保存图片
            with open(img_path, 'wb') as img_file:
                img_file.write(img)
            # 获取图片所在的段落位置
            for i, para in enumerate(doc.paragraphs):
                if rel.target_ref in para._element.xml:
                    image_positions.append((i, img_name))
                    break
    return image_positions

def extract_tables(docx_path, excel_folder):
    """提取表格并保存为Excel文件，并返回表格位置字典"""
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)
    doc = docx.Document(docx_path)
    table_positions = []
    for i, table in enumerate(doc.tables):
        df = pd.DataFrame(columns=[f'Column{j+1}' for j in range(len(table.columns))])
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            if len(row_data) != len(df.columns):
                print(f"Skipping row with mismatched columns: {row_data}")
                continue
            df.loc[len(df)] = row_data
        excel_path = os.path.join(excel_folder, f'table_{i+1}.xlsx')
        df.to_excel(excel_path, index=False)
        
        # 获取表格所在的段落位置
        for j, para in enumerate(doc.paragraphs):
            if table._element in para._element.getparent().getchildren():
                table_positions.append((j, f'table_{i+1}.xlsx'))
                break
    return table_positions

def convert_excel_to_markdown(excel_path):
    """将Excel表格转换为Markdown格式"""
    df = pd.read_excel(excel_path)
    return df.to_markdown(index=False)

def extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions):
    """提取文本并转换为Markdown格式，并在图片和表格位置添加占位符"""
    doc = docx.Document(docx_path)
    markdown_text = ''
    
    # 遍历文档段落
    for i, para in enumerate(doc.paragraphs):
        # 检查是否需要在当前段落插入图片占位符
        for img_index, img_name in sorted(image_positions, key=lambda x: x[0]):
            if img_index == i:
                img_path = os.path.join(image_folder, img_name)
                markdown_text += f'![{img_name}]({img_path})\n\n'
                image_positions.remove((img_index, img_name))  # 移除已处理的图片

        # 检查是否需要在当前段落插入表格占位符
        for table_index, table_file in sorted(table_positions, key=lambda x: x[0]):
            if table_index == i:
                excel_path = os.path.join(excel_folder, table_file)
                table_md = convert_excel_to_markdown(excel_path)
                markdown_text += f'\n\n{table_md}\n\n'
                table_positions.remove((table_index, table_file))  # 移除已处理的表格
        
        # 处理段落内容
        if para.style and para.style.name.startswith('Heading'):
            level = para.style.name.split('Heading')[1] if len(para.style.name.split('Heading')) > 1 else "1"
            markdown_text += f'{"#" * int(level)} {para.text}\n\n'
        elif para.style and para.style.name == 'Normal':
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            markdown_text += f'{para.text}\n\n'
        elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            markdown_text += f':::center\n{para.text}\n:::\n\n'
        
        # 处理文本格式
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f'**{text}**'
            if run.italic:
                text = f'*{text}*'
            if run.underline:
                text = f'_{text}_'
            if 'http' in run.text:
                text = f'[{text}]({text})'
            markdown_text += text

    # 保存Markdown文件
    markdown_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.md'
    markdown_path = os.path.join(markdown_folder, markdown_filename)
    with open(markdown_path, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_text)


def process_docx_folder(docx_folder, output_folder):
    """处理文件夹中的所有docx文件并生成Markdown文件"""
    image_folder = os.path.join(output_folder, 'images')
    excel_folder = os.path.join(output_folder, 'excel')
    markdown_folder = os.path.join(output_folder, 'markdown')
    
    if not os.path.exists(markdown_folder):
        os.makedirs(markdown_folder)
    
    for filename in os.listdir(docx_folder):
        if filename.endswith('.docx'):
            docx_path = os.path.join(docx_folder, filename)
            # 提取图片和表格，并获取它们的位置
            image_positions = extract_images(docx_path, image_folder)
            table_positions = extract_tables(docx_path, excel_folder)
            # 生成Markdown文件
            extract_text_to_markdown(docx_path, markdown_folder, image_folder, excel_folder, image_positions, table_positions)

# 使用示例
docx_folder = r""  # 输入的.docx文件夹路径
output_folder = r""  # 输出的Markdown文件夹路径

process_docx_folder(docx_folder, output_folder)
