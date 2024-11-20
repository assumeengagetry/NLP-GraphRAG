import os
import shutil
import logging
from tqdm import tqdm
import pandas as pd
from docx import Document
import subprocess

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def move_docx_files(src_folder, dest_folder):
    for filename in os.listdir(src_folder):
        if filename.endswith(".docx"):
            src_file = os.path.join(src_folder, filename)
            dest_file = os.path.join(dest_folder, filename)
            shutil.move(src_file, dest_file)
            logging.info(f"Moved '{src_file}' to '{dest_file}'")

def extract_text(docx_file, output_md_file):
    subprocess.run(["pandoc", docx_file, "-o", output_md_file], check=True)
    logging.info(f"Converted '{docx_file}' to '{output_md_file}'")

def extract_tables(docx_file):
    doc = Document(docx_file)
    tables_data = []
    for table in doc.tables:
        df = pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows])
        tables_data.append(df)
    return tables_data

def save_tables_to_excel(tables_data, excel_file):
    if not tables_data:
        logging.info(f"No tables to convert for {excel_file}")
        return
    with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
        for i, table_df in enumerate(tables_data):
            if not table_df.empty:
                table_df.to_excel(writer, sheet_name=f'Table {i+1}', index=False)
            else:
                logging.info(f"Table {i+1} is empty, skipped.")
    logging.info(f"Tables saved to {excel_file}")

def excel_to_markdown(excel_file, output_folder):
    xls = pd.ExcelFile(excel_file)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        markdown_table = df.to_markdown(index=False)
        markdown_file = os.path.join(output_folder, f"{os.path.splitext(os.path.basename(excel_file))[0]}_{sheet_name}.md")
        with open(markdown_file, 'w', encoding='utf-8') as md_file:
            md_file.write(f"## {sheet_name}\n\n{markdown_table}\n\n")
        logging.info(f"Markdown file saved: {markdown_file}")

def extract_images(docx_file, image_folder):
    doc = Document(docx_file)
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part.blob
            image_filename = os.path.basename(rel.target_ref)
            image_path = os.path.join(image_folder, image_filename)
            with open(image_path, 'wb') as img_file:
                img_file.write(image)
    logging.info(f"Images extracted to {image_folder}")

def process_docx_file(docx_file, output_md_folder, output_excel_folder, output_image_folder, output_table_md_folder):
    if not os.path.exists(output_md_folder):
        os.makedirs(output_md_folder)
    if not os.path.exists(output_excel_folder):
        os.makedirs(output_excel_folder)
    if not os.path.exists(output_image_folder):
        os.makedirs(output_image_folder)
    if not os.path.exists(output_table_md_folder):
        os.makedirs(output_table_md_folder)

    output_md_file = os.path.join(output_md_folder, os.path.splitext(os.path.basename(docx_file))[0] + '.md')
    extract_text(docx_file, output_md_file)

    tables_data = extract_tables(docx_file)
    excel_file = os.path.join(output_excel_folder, os.path.splitext(os.path.basename(docx_file))[0] + '.xlsx')
    save_tables_to_excel(tables_data, excel_file)
    if os.path.exists(excel_file):
        excel_to_markdown(excel_file, output_table_md_folder)

    extract_images(docx_file, output_image_folder)

def convert_doc_to_docx(input_directory, output_directory, libreoffice_path):
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    for filename in os.listdir(input_directory):
        if filename.endswith(".doc"):
            input_path = os.path.join(input_directory, filename)
            output_path = os.path.join(output_directory, os.path.splitext(filename)[0] + '.docx')
            command = [libreoffice_path, "--headless", "--convert-to", "docx", "--outdir", output_directory, input_path]
            subprocess.run(command, check=True)
            print(f"Converted {input_path} to {output_path}")

def main():
    src_folder = r"C:\Users\HP\Desktop\大创\input"
    dest_folder = r"C:\Users\HP\Desktop\大创\output_step1"
    convert_doc_to_docx(src_folder, dest_folder, "C:\\Program Files\\LibreOffice\\program\\soffice.com")
    
    input_folder = dest_folder
    output_md_folder = r"C:\Users\HP\Desktop\大创\output_step2"
    output_excel_folder = r"C:\Users\HP\Desktop\大创\chart_excel"
    output_image_folder = r"C:\Users\HP\Desktop\大创\pictures"
    output_table_md_folder = r"C:\Users\HP\Desktop\大创\chart_child"
    
    for filename in tqdm(os.listdir(input_folder), desc="Processing files"):
        if filename.endswith(".docx"):
            docx_file_path = os.path.join(input_folder, filename)
            process_docx_file(docx_file_path, output_md_folder, output_excel_folder, output_image_folder, output_table_md_folder)

if __name__ == "__main__":
    main()