import os
from docx import Document
import pandas as pd

def extract_text(docx_file):
    doc = Document(docx_file)
    markdown_content = []
    for para in doc.paragraphs:
        markdown_content.append(para.text + '\n\n')
    return ''.join(markdown_content)

def extract_tables(docx_file):
    doc = Document(docx_file)
    tables_data = []
    for table in doc.tables:
        df = pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows])
        tables_data.append(df)
    return tables_data

def save_tables_to_excel(tables_data, excel_file):
    if not tables_data:
        print("No tables to convert.")
        return
    with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
        for i, table_df in enumerate(tables_data):
            table_df.to_excel(writer, sheet_name=f'Table {i+1}', index=False)

def excel_to_markdown(excel_file):
    xls = pd.ExcelFile(excel_file)
    markdown_content = ""
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        markdown_table = df.to_markdown(index=False)
        markdown_content += f"## {sheet_name}\n\n{markdown_table}\n\n"
    return markdown_content

def extract_images(docx_file, image_folder):
    doc = Document(docx_file)
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part.blob
            image_filename = os.path.basename(rel.target_ref)
            image_path = os.path.join(image_folder, image_filename)
            with open(image_path, 'wb') as img_file:
                img_file.write(image)
    print(f"Images extracted to {image_folder}")

def main(docx_file, output_md_folder, output_excel_folder, output_image_folder):
    # Ensure output directories exist
    if not os.path.exists(output_md_folder):
        os.makedirs(output_md_folder)
    if not os.path.exists(output_excel_folder):
        os.makedirs(output_excel_folder)
    if not os.path.exists(output_image_folder):
        os.makedirs(output_image_folder)

    # Extract text content
    markdown_content = extract_text(docx_file)
    
    # Extract tables data
    tables_data = extract_tables(docx_file)
    
    # Save tables as Excel file
    excel_file = os.path.splitext(os.path.basename(docx_file))[0] + '.xlsx'
    if tables_data:
        save_tables_to_excel(tables_data, os.path.join(output_excel_folder, excel_file))
        markdown_tables_content = excel_to_markdown(os.path.join(output_excel_folder, excel_file))
    else:
        markdown_tables_content = ""

    # Combine text and tables Markdown content
    markdown_content_with_tables = markdown_content + "\n" + markdown_tables_content
    
    # Save Markdown file
    markdown_file = os.path.splitext(os.path.basename(docx_file))[0] + '.md'
    with open(os.path.join(output_md_folder, markdown_file), 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_content_with_tables)
    print(f"Markdown file saved: {os.path.join(output_md_folder, markdown_file)}")

    # Extract images
    extract_images(docx_file, output_image_folder)

if __name__ == "__main__":
    input_folder = r"C:\Users\HP\Desktop\大创\output_step1"  # 输入文件夹路径
    output_md_folder = r"C:\Users\HP\Desktop\大创\output_step2_md"  # 输出Markdown文件夹路径
    output_excel_folder = r"C:\Users\HP\Desktop\大创\chart_excel"  # 输出Excel文件夹路径
    output_image_folder = r"C:\Users\HP\Desktop\大创\pictures"  # 输出图片文件夹路径

    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            docx_file_path = os.path.join(input_folder, filename)
            main(docx_file_path, output_md_folder, output_excel_folder, output_image_folder)